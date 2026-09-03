"""Local video -> audio transcript DOCX + subtitle-aligned screenshot PDF."""
from __future__ import annotations

from dataclasses import dataclass, asdict, field
from pathlib import Path
from datetime import datetime
from collections import Counter
import json
import os
import re
import shutil
import subprocess
import tempfile
import threading
import unicodedata
import wave

import numpy as np
import pysubs2
from PIL import Image
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.lib.utils import ImageReader
from compat import process_options,ocr_command,folder_name
from notices import SHORT_NOTICE,VERSION as NOTICE_VERSION,export_notice

ROOT = Path(__file__).resolve().parent
RATE = 16000


class Cancelled(Exception):
    pass


@dataclass
class Segment:
    start: float
    end: float
    text: str
    capture: float | None = None
    reason: str = "字幕 / 语音分段"


@dataclass
class Options:
    video: str
    output: str
    subtitle: str = ""
    mode: str = "auto"  # auto / ocr / speech
    word_source: str = "audio"
    language: str = "auto"
    roi: tuple = (0.0, 0.72, 1.0, 0.26)
    interval: float = 0.5
    per_page: int = 2
    start: float = 0.0
    end: float = 0.0
    visual_changes: bool = True
    visual_sensitivity: str = "normal"
    visual_gap: float = 0.4
    evidence: dict = field(default_factory=dict)
    attachments: list = field(default_factory=list)
    result_name: str = ""


def timestamp(seconds: float) -> str:
    ms = max(0, round(seconds * 1000))
    sec, ms = divmod(ms, 1000)
    minute, sec = divmod(sec, 60)
    hour, minute = divmod(minute, 60)
    return f"{hour:02}:{minute:02}:{sec:02}.{ms:03}"


def executable(name: str) -> str:
    binary=name+".exe" if os.name=="nt" else name
    for path in (ROOT / "bin" / binary, Path("/opt/homebrew/bin") / binary, Path("/usr/local/bin") / binary):
        if path.is_file():
            return str(path)
    result = shutil.which(name)
    if not result:
        raise RuntimeError(f"没有找到 {name}，请重新运行安装程序。")
    return result


def run(args, stop=None):
    """Drain pipes while checking cancellation, including during long ffmpeg calls."""
    with subprocess.Popen([str(a) for a in args], stdout=subprocess.PIPE, stderr=subprocess.PIPE,**process_options()) as proc:
        try:
            while True:
                if stop and stop.is_set():
                    raise Cancelled()
                try:
                    out, err = proc.communicate(timeout=0.2)
                    break
                except subprocess.TimeoutExpired:
                    pass
            if proc.returncode:
                raise RuntimeError(err.decode("utf-8", "replace")[-1800:] or "外部程序运行失败")
            return out
        finally:
            if proc.poll() is None:
                proc.kill()
                proc.wait()


def probe(path, stop=None):
    data = json.loads(run([executable("ffprobe"), "-v", "error", "-show_format", "-show_streams", "-of", "json", path], stop))
    video = next((s for s in data["streams"] if s["codec_type"] == "video" and not s.get("disposition", {}).get("attached_pic")), None)
    if video is None:
        raise ValueError("这个文件没有可读取的视频画面，请选择视频文件。")
    duration = float(data.get("format", {}).get("duration") or video.get("duration") or 0)
    if duration <= 0:
        raise ValueError("无法读取视频时长，请尝试将视频导出为 MP4 后再导入。")
    return data, duration


def frame(video, at, output, width=1600, stop=None):
    run([executable("ffmpeg"), "-v", "error", "-nostdin", "-y", "-ss", f"{max(0, at):.6f}", "-i", video,
         "-map", "0:v:0", "-frames:v", "1", "-vf", f"scale='min({width},iw)':-2", "-q:v", "2", output], stop)
    if not Path(output).is_file():
        raise RuntimeError(f"未能读取 {timestamp(at)} 处的视频画面。")


def clip_segments(rows, start, end):
    return [Segment(max(start, r.start), min(end, r.end), r.text,
                    r.capture if r.capture is not None and start <= r.capture < end else None)
            for r in rows if r.text.strip() and r.end > start and r.start < end]


def load_subtitles(path, start, end):
    try:
        subs = pysubs2.load(str(path), encoding="utf-8-sig")
    except UnicodeDecodeError:
        subs = pysubs2.load(str(path), encoding="gb18030")
    rows = [Segment(s.start / 1000, s.end / 1000, s.plaintext.strip()) for s in subs if not s.is_comment and s.end > s.start]
    # ASS and bilingual files may contain simultaneous cues. Capture each visible
    # combination once, rather than creating conflicting pages at the same time.
    events = {}
    for i, row in enumerate(rows):
        events.setdefault(row.start, [[], []])[1].append(i)
        events.setdefault(row.end, [[], []])[0].append(i)
    active, result, previous = set(), [], None
    for time in sorted(events):
        if previous is not None and active and time > previous:
            text = "\n".join(dict.fromkeys(rows[i].text for i in sorted(active) if rows[i].text))
            if result and result[-1].text == text and abs(result[-1].end - previous) < 1e-6:
                result[-1].end = time
            else:
                result.append(Segment(previous, time, text))
        remove, add = events[time]
        active.difference_update(remove)
        active.update(add)
        previous = time
    return clip_segments(result, start, end)


def normalize(text):
    return "".join(c.lower() for c in unicodedata.normalize("NFKC", text) if c.isalnum())


def group_observations(observations, step, end):
    """Retain every detected text change; do not fuzzy-merge changed numbers/words."""
    rows = []
    active = None
    variants = Counter()
    for time, text in observations:
        key = normalize(text)
        if active and key != normalize(active.text):
            active.end = time
            active.text = variants.most_common(1)[0][0]
            rows.append(active)
            active, variants = None, Counter()
        if key and active is None:
            active = Segment(time, min(end, time + step), text, time)
        if active:
            active.end = min(end, time + step)
            variants[text] += 1
    if active:
        active.text = variants.most_common(1)[0][0]
        rows.append(active)
    return rows


def ocr_scan(opt, start, end, work, stop, progress):
    # Decode sequentially into one reusable image. Long recordings never create
    # an intermediate JPEG for every sampled frame or accumulate raw video RAM.
    roi = opt.roi
    if not (len(roi) == 4 and all(0 <= x <= 1 for x in roi) and roi[2] > 0 and roi[3] > 0
            and roi[0] + roi[2] <= 1.00001 and roi[1] + roi[3] <= 1.00001):
        raise ValueError("字幕框必须位于视频画面内。")
    sample = work / "ocr-frame.bmp"
    initial = work / "first.jpg"
    frame(opt.video, start, initial, width=1280, stop=stop)
    with Image.open(initial) as image:
        w, h = image.size
    cmd = [executable("ffmpeg"), "-v", "error", "-nostdin", "-ss", str(start), "-i", opt.video,
           "-t", str(end-start), "-map", "0:v:0", "-vf", f"fps=fps=1/{opt.interval}:start_time=0:round=up,scale={w}:{h}",
           "-f", "rawvideo", "-pix_fmt", "rgb24", "pipe:1"]
    observations = []
    with (work / "decode.log").open("wb") as log, subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=log,**process_options()) as decoder, \
            subprocess.Popen(ocr_command(), stdin=subprocess.PIPE, stdout=subprocess.PIPE,
                             stderr=log, text=True,encoding="utf-8", bufsize=1,**process_options()) as ocr:
        # Cancellation can interrupt blocking reads and Vision requests as well.
        finished = threading.Event()
        def cancel_watch():
            while not finished.wait(0.15):
                if stop.is_set():
                    for p in (decoder, ocr):
                        try:
                            p.kill()
                        except ProcessLookupError:
                            pass
                    return
        watcher = threading.Thread(target=cancel_watch, daemon=True)
        watcher.start()
        try:
            index = 0
            size = w * h * 3
            while True:
                if stop.is_set():
                    raise Cancelled()
                data = decoder.stdout.read(size)
                if not data:
                    break
                if len(data) != size:
                    raise RuntimeError("视频解码提前中断。")
                time = start + index * opt.interval
                if time >= end:
                    break
                Image.frombytes("RGB", (w, h), data).save(sample)
                ocr.stdin.write(json.dumps({"path": str(sample), "roi": roi}) + "\n")
                ocr.stdin.flush()
                response = ocr.stdout.readline()
                if stop.is_set():
                    raise Cancelled()
                if not response:
                    raise RuntimeError("字幕识别服务退出，请重新打开程序。")
                result = json.loads(response)
                if "error" in result:
                    raise RuntimeError("字幕识别失败：" + result["error"])
                observations.append((time, result["text"]))
                index += 1
                progress(8 + 40 * (time-start)/(end-start), f"识别画面字幕 · {timestamp(time)} / {timestamp(end)}")
            if stop.is_set():
                raise Cancelled()
            if decoder.poll() not in (None, 0):
                raise RuntimeError("视频解码失败，请检查文件是否完整。")
        finally:
            finished.set()
            for p in (decoder, ocr):
                if p.poll() is None:
                    p.kill()
                p.wait()
            watcher.join(timeout=1)
    return group_observations(observations, opt.interval, end)


def transcribe(opt, start, end, work, stop, progress):
    import sherpa_onnx
    progress(49, "提取音轨，准备本地语音识别…")
    audio = work / "audio.wav"
    run([executable("ffmpeg"), "-v", "error", "-nostdin", "-y", "-ss", str(start), "-i", opt.video,
         "-t", str(end-start), "-map", "0:a:0", "-vn", "-ac", "1", "-ar", str(RATE), "-c:a", "pcm_s16le", audio], stop)
    recognizer = sherpa_onnx.OfflineRecognizer.from_sense_voice(
        model=str(ROOT / "models/sensevoice.int8.onnx"), tokens=str(ROOT / "models/tokens.txt"),
        num_threads=4, language="" if opt.language == "auto" else opt.language, use_itn=True)
    config = sherpa_onnx.VadModelConfig()
    config.silero_vad.model = str(ROOT / "models/silero_vad.onnx")
    config.silero_vad.threshold = 0.5
    config.silero_vad.min_silence_duration = 0.35
    config.silero_vad.min_speech_duration = 0.15
    config.silero_vad.max_speech_duration = 8.0
    config.sample_rate = RATE
    vad = sherpa_onnx.VoiceActivityDetector(config, buffer_size_in_seconds=30)
    rows = []
    def drain():
        while not vad.empty():
            if stop.is_set():
                raise Cancelled()
            segment = vad.front
            begin = start + segment.start/RATE
            samples = np.asarray(segment.samples, dtype=np.float32)
            finish = min(end, begin + len(samples)/RATE)
            stream = recognizer.create_stream()
            stream.accept_waveform(RATE, samples)
            recognizer.decode_stream(stream)
            text = re.sub(r"<\|.*?\|>", "", stream.result.text).strip()
            if text and finish > begin:
                rows.append(Segment(begin, finish, text))
            vad.pop()
    with wave.open(str(audio), "rb") as reader:
        total = reader.getnframes()
        done = 0
        while True:
            if stop.is_set():
                raise Cancelled()
            data = reader.readframes(512)
            if not data:
                break
            samples = np.frombuffer(data, dtype="<i2").astype(np.float32) / 32768.0
            done += len(samples)
            if len(samples) < 512:
                samples = np.pad(samples, (0, 512-len(samples)))
            vad.accept_waveform(samples)
            drain()
            if done % (512*32) == 0:
                progress(50 + 24 * done/max(1,total), f"语音转文字 · {timestamp(start + done/RATE)} / {timestamp(end)}")
        vad.flush()
        drain()
    audio.unlink()
    return rows


def write_srt(rows, path):
    subs = pysubs2.SSAFile()
    subs.events = [pysubs2.SSAEvent(start=round(r.start*1000), end=round(r.end*1000), text=r.text.replace("\n", r"\N")) for r in rows]
    subs.save(str(path), format_="srt", encoding="utf-8")


def write_docx(title, rows, path, source, start, end):
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "PingFang SC")
    document.add_heading(title, 0)
    document.add_paragraph(f"文字来源：{source} ｜ 视频范围：{timestamp(start)}–{timestamp(end)}")
    document.add_paragraph("自动识别可能有误；时间为视频中的原始位置。")
    document.add_paragraph(SHORT_NOTICE)
    if not rows:
        document.add_paragraph("本次没有识别到可用的语音文字。")
    for row in rows:
        p = document.add_paragraph()
        r = p.add_run(f"{timestamp(row.start)} – {timestamp(row.end)}")
        r.bold = True
        p.paragraph_format.keep_with_next = True
        document.add_paragraph(row.text)
    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = 2
    footer.add_run("第 ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)
    footer.add_run(" 页")
    document.save(path)


def wrap_text(text, font, size, width):
    lines = []
    for paragraph in text.splitlines() or [""]:
        current = ""
        for char in paragraph:
            if current and pdfmetrics.stringWidth(current+char, font, size) > width:
                lines.append(current)
                current = ""
            current += char
        lines.append(current)
    return lines


def write_pdf(title, rows, images, path, per_page, source, stop):
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    font = "STSong-Light"
    width, height = (842, 595) if per_page == 1 else (595, 842)
    c = canvas.Canvas(str(path), pagesize=(width, height), pageCompression=1)
    c.setTitle(title + " · 字幕截图")
    c.setAuthor("视频成册")
    margin = 30
    slot = (height - 86)/per_page
    pages = (len(rows)+per_page-1)//per_page
    for i, (row, image_path) in enumerate(zip(rows, images)):
        if stop.is_set():
            raise Cancelled()
        position = i % per_page
        if position == 0:
            c.setFillColorRGB(.1, .2, .24)
            c.setFont(font, 13)
            display_title = title
            while pdfmetrics.stringWidth(display_title, font, 13) > width-2*margin:
                display_title = display_title[:-2] + "…"
            c.drawString(margin, height-29, display_title)
            c.setFont(font, 8)
            c.setFillColorRGB(.42, .47, .5)
            c.drawString(margin, 20, source)
            c.setFont(font,7)
            c.drawString(margin,9,"自动生成、需人工核对；不提供公证或可信时间戳。")
            c.setFont(font,8)
            c.drawRightString(width-margin, 20, f"{i//per_page+1} / {pages}")
        top = height-48-position*slot
        c.setFont(font, 9)
        c.setFillColorRGB(.2, .35, .38)
        at = row.capture if row.capture is not None else (row.start+row.end)/2
        c.drawString(margin, top-11, f"{i+1:04d}   {timestamp(at)}   {row.reason}")
        text_size = 10
        lines = wrap_text(row.text, font, text_size, width-2*margin)
        # Keep captions readable; full long cues remain available in the SRT.
        shown = lines[:6]
        if len(lines) > 6:
            shown[-1] = shown[-1][:max(0, len(shown[-1])-16)] + " …（全文见截图对应字幕.srt）"
        caption_h = len(shown)*14+14
        box_top = top-22
        box_h = slot-35-caption_h
        with Image.open(image_path) as im:
            iw, ih = im.size
            scale = min((width-2*margin)/iw, box_h/ih)
            dw, dh = iw*scale, ih*scale
            c.drawImage(ImageReader(im), margin, box_top-dh, dw, dh)
        c.setFillColorRGB(.12, .17, .20)
        c.setFont(font, text_size)
        y = box_top-dh-17
        for line in shown:
            c.drawString(margin, y, line)
            y -= 14
        if position == per_page-1 or i == len(rows)-1:
            c.showPage()
    c.save()


def process(opt: Options, stop=None, progress=lambda p, s: None):
    import evidence
    import visual
    started_at = evidence.local_time()
    stop = stop or threading.Event()
    video = Path(opt.video).expanduser().resolve()
    opt.video = str(video)
    if not video.is_file():
        raise ValueError("视频文件不存在。")
    initial_stat=(video.stat().st_size,video.stat().st_mtime_ns)
    opt.evidence=dict(opt.evidence)
    opt.evidence["evidence_id"]=opt.evidence.get("evidence_id") or evidence.new_evidence_id()
    if opt.interval < 0.1 or opt.per_page not in (1, 2):
        raise ValueError("无效的采样间隔或 PDF 布局。")
    if opt.visual_sensitivity not in {"high","normal","low"} or opt.visual_gap < 0:
        raise ValueError("无效的画面检测设置。")
    def check():
        if stop.is_set(): raise Cancelled()
    progress(1, "读取视频信息…")
    meta, duration = probe(video, stop)
    start = max(0, opt.start)
    end = min(duration, opt.end if opt.end > 0 else duration)
    if end <= start:
        raise ValueError("结束时间必须晚于开始时间，而且范围必须位于视频内。")
    output = Path(opt.output).expanduser().resolve()
    output.mkdir(parents=True, exist_ok=True)
    safe_title = video.stem[:90]
    document_title=f"{opt.evidence['evidence_id']} / {safe_title}"
    name = folder_name(opt.result_name) or f"{safe_title}_{datetime.now():%Y%m%d_%H%M%S_%f}"
    destination = output / name
    if destination.exists():
        raise ValueError("本次结果文件夹已存在，请更换名称；已有结果不会被覆盖。")
    warnings = []
    with tempfile.TemporaryDirectory(prefix=".视频成册-处理中-", dir=output) as tmp:
        work = Path(tmp)
        package = work / "result"
        package.mkdir()
        rows, source, speech = [], "", None
        if opt.subtitle:
            rows = load_subtitles(opt.subtitle, start, end)
            source = "字幕文件"
            if not rows:
                raise ValueError("所选字幕文件在这个时间范围内没有字幕。")
        elif opt.mode == "auto":
            streams = [s for s in meta["streams"] if s["codec_type"] == "subtitle" and s.get("codec_name") in
                       {"subrip", "ass", "ssa", "mov_text", "webvtt", "text"}]
            streams.sort(key=lambda s: (not s.get("disposition", {}).get("default", 0), s["index"]))
            if streams:
                progress(5, "提取视频内置字幕…")
                srt = work / "embedded.srt"
                run([executable("ffmpeg"), "-v", "error", "-nostdin", "-y", "-i", video,
                     "-map", f"0:{streams[0]['index']}", "-c:s", "srt", srt], stop)
                rows = load_subtitles(srt, start, end)
                source = "视频内置字幕"
        if not rows and opt.mode != "speech":
            progress(8, "逐段检查画面字幕…")
            rows = ocr_scan(opt, start, end, work, stop, progress)
            source = "画面字幕 OCR"
            warnings.append(f"画面字幕每 {opt.interval:g} 秒检查一次，短于该间隔的字幕可能遗漏；文字识别可能出现错字或重复。")
        has_audio = any(s["codec_type"] == "audio" for s in meta["streams"])
        if has_audio and (opt.word_source == "audio" or not rows):
            speech = transcribe(opt, start, end, work, stop, progress)
        elif opt.word_source == "audio":
            warnings.append("视频没有音轨，Word 改用字幕文字。")
        if not rows:
            if speech:
                rows = speech
                source = "语音分段（非原视频字幕切换）"
                warnings.append("未使用可识别的画面字幕，已按语音停顿分段截图；分段时间不等同于原视频字幕时间。")
            elif not opt.visual_changes:
                raise ValueError("没有找到可用字幕或语音。请调整字幕框、导入 SRT，或选择有声音的视频。")
        subtitle_rows = rows
        word_rows = speech if opt.word_source == "audio" and speech else subtitle_rows
        word_source = "音频自动转写" if opt.word_source == "audio" and speech else source
        if not word_rows:
            word_source = "无可用语音或字幕"
        if opt.visual_changes:
            progress(74,"检查人物动作、插入画面和其他视觉变化…")
            times = visual.scan(opt,start,end,work,stop,progress)
            rows = visual.combine(subtitle_rows,times,start,end,opt.interval)
            source = (source+" + 画面变化") if subtitle_rows else "画面变化（无字幕）"
            warnings.append(f"画面变化按 {opt.interval:g} 秒采样检测，可能漏掉更短片段；灵敏度 {opt.visual_sensitivity}，常规截图最小间隔 {opt.visual_gap:g} 秒。不是人物身份或动作含义识别。")
        if not rows:
            raise ValueError("所选时间范围没有读取到可保存画面。")
        if opt.word_source == "audio" and speech == []:
            warnings.append("音轨中未识别到语音，Word 改用字幕文字。")
        progress(76, f"生成 Word，准备保存 {len(rows)} 张截图…")
        write_docx(document_title, word_rows, package / "文字稿.docx", word_source, start, end)
        (package / "文字稿.txt").write_text("\n\n".join(f"[{timestamp(r.start)} – {timestamp(r.end)}]\n{r.text}" for r in word_rows), encoding="utf-8")
        write_srt(subtitle_rows, package / "截图对应字幕.srt")
        if speech is not None:
            write_srt(speech, package / "音频转写.srt")
        image_dir = package / "截图"
        image_dir.mkdir()
        images = []
        for i, row in enumerate(rows):
            if stop.is_set():
                raise Cancelled()
            # OCR captures the exact sampled frame that actually contained text.
            # Timed subtitle files capture the middle, away from transitions.
            at = row.capture if row.capture is not None else (row.start+row.end)/2
            path = image_dir / f"{i+1:05d}_{timestamp(at).replace(':', '-')}.jpg"
            frame(video, min(at, end-.001), path, stop=stop)
            images.append(path)
            progress(77+19*(i+1)/len(rows), f"保存截图 · {i+1} / {len(rows)}")
        progress(97, "排版并生成 PDF…")
        write_pdf(document_title, rows, images, package / "字幕截图.pdf", opt.per_page, source, stop)
        progress(98,"记录原文件校验值和取证信息…")
        original_hash = evidence.fingerprint(video,check)
        if (video.stat().st_size,video.stat().st_mtime_ns)!=initial_stat:
            raise RuntimeError("原视频在处理期间被修改，请等待文件保存完成后重试。")
        attachment_records=[]
        for i,attachment in enumerate(opt.attachments,1):
            check()
            original=Path(attachment)
            attachment_dir=package/"评论区等附件"
            attachment_dir.mkdir(exist_ok=True)
            target=attachment_dir/f"{i:03d}_{original.name}"
            shutil.copy2(original,target)
            attachment_records.append({"file":str(target.relative_to(package)),"original":str(original),"sha256":evidence.fingerprint(target,check)})
        record=evidence.make_record(opt.evidence, export_started_at=started_at,export_finished_at=evidence.local_time(),
            original_file=str(video),sha256=original_hash,range=f"{timestamp(start)} – {timestamp(end)}",
            output_folder=str(destination),generated_files="文字稿.docx\n字幕截图.pdf\n截图/\n取证信息.xlsx\n取证信息.json\n截图对应字幕.srt\n免责声明与使用边界.txt",
            attachment_names="\n".join(x["file"] for x in attachment_records),screenshot_count=len(rows))
        evidence.write_record(record,package)
        export_notice(package)
        report = {"video": str(video), "range": [start,end], "screenshot_source": source, "word_source": word_source,
                  "screenshots": len(rows), "transcript_segments": len(word_rows), "warnings": warnings,
                  "options": asdict(opt), "segments": [asdict(r) for r in rows],"evidence":record,"attachments":attachment_records,"disclaimer_version":NOTICE_VERSION}
        (package / "处理记录.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        (package / "阅读说明.txt").write_text(
            f"视频成册\n\n截图依据：{source}\nWord 来源：{word_source}\n截图数量：{len(rows)}\n"
            "截图保留完整视频画面，PDF 图下附对应文字。\n时间戳均对应原视频。\n"
            "超过 PDF 图下注释空间的长字幕，完整内容保存在“截图对应字幕.srt”。\n\n" + SHORT_NOTICE + "\n完整声明见“免责声明与使用边界.txt”。\n\n" + "\n".join(warnings), encoding="utf-8")
        if stop.is_set():
            raise Cancelled()
        # Reserve the chosen name before moving the completed files. Never replace
        # a directory created by another task after the initial existence check.
        destination.mkdir()
        try:
            for child in package.iterdir(): child.rename(destination/child.name)
        except Exception:
            # Completed partial output is retained if the filesystem move fails.
            raise RuntimeError(f"移动结果时出错，已移动的文件保留在：{destination}")
    try:
        ledger=evidence.append_ledger(record,output)
        report["ledger"]=str(ledger)
    except Exception as error:
        warnings.append(f"Word、PDF 和单条取证信息已保存，但累计台账未追加：{error}")
        report["ledger_error"]=str(error)
    (destination/"处理记录.json").write_text(json.dumps(report,ensure_ascii=False,indent=2),encoding="utf-8")
    progress(100, f"完成 · {len(rows)} 张截图，Word 和 PDF 已保存")
    return destination, report
